from docx import Document
from io import BytesIO

from docx.table import Table
from docx.text.paragraph import Paragraph

from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------
# ITERATE BLOCKS
# ---------------------------------------------------

def iter_block_items(parent):

    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    parent_elm = parent.element.body

    for child in parent_elm.iterchildren():

        if isinstance(child, CT_P):
            yield Paragraph(child, parent)

        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# ---------------------------------------------------
# CLEAN TEXT
# ---------------------------------------------------

def clean_text(text):

    return (
        str(text)
        .replace("\n", "")
        .replace(" ", "")
        .strip()
    )


# ---------------------------------------------------
# MAIN FUNCTION
# ---------------------------------------------------

def process_document(file, account_number):

    source_doc = Document(file)

    output_doc = Document()

    # ---------------------------------------------------
    # PAGE SETTINGS
    # ---------------------------------------------------

    section = output_doc.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE

    new_width = section.page_height
    new_height = section.page_width

    section.page_width = new_width
    section.page_height = new_height

    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    # ---------------------------------------------------
    # HEADING
    # ---------------------------------------------------

    heading = output_doc.add_paragraph()

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = heading.add_run(
        "Format for restoration / rejection of unauthorized amount"
    )

    run.bold = True
    run.font.size = Pt(14)

    annexure = output_doc.add_paragraph()

    annexure.alignment = WD_ALIGN_PARAGRAPH.CENTER

    annexure_run = annexure.add_run(
        "Annexure-B"
    )

    annexure_run.bold = True
    annexure_run.font.size = Pt(12)

    output_doc.add_paragraph()

    # ---------------------------------------------------
    # VARIABLES
    # ---------------------------------------------------

    started_copying = False

    matched_channel = None

    inside_correct_channel = False

    # ---------------------------------------------------
    # ITERATE SOURCE DOCUMENT
    # ---------------------------------------------------

    for block in iter_block_items(source_doc):

        # ---------------------------------------------------
        # TABLES
        # ---------------------------------------------------

        if isinstance(block, Table):

            try:

                rows = block.rows

                if len(rows) < 2:
                    continue

                headers = [
                    cell.text.strip()
                    for cell in rows[0].cells
                ]

                values = [
                    cell.text.strip()
                    for cell in rows[1].cells
                ]

                matched = False

                for value in values:

                    if clean_text(value) == clean_text(account_number):

                        matched = True
                        break

                # ---------------------------------------------------
                # MATCH FOUND
                # ---------------------------------------------------

                if (
                    "Account Number" in headers
                    and matched
                ):

                    started_copying = True

                    matched_channel = (
                        values[5]
                        .strip()
                        .upper()
                    )

                    # ---------------------------------------------
                    # CREATE TABLE
                    # ---------------------------------------------

                    num_rows = len(rows)
                    num_cols = len(rows[0].cells)

                    final_table = output_doc.add_table(
                        rows=num_rows,
                        cols=num_cols
                    )

                    final_table.style = "Table Grid"

                    final_table.alignment = 1
                    final_table.autofit = True

                    for i, row in enumerate(rows):

                        for j, cell in enumerate(row.cells):

                            new_cell = final_table.cell(i, j)

                            new_cell.text = cell.text.strip()

                            for para in new_cell.paragraphs:

                                para.alignment = (
                                    WD_ALIGN_PARAGRAPH.CENTER
                                )

                                for run in para.runs:

                                    run.font.size = Pt(7)

                    spacing = output_doc.add_paragraph()

                    spacing.paragraph_format.space_after = Pt(12)

                    continue

                # ---------------------------------------------------
                # STOP NEXT CUSTOMER
                # ---------------------------------------------------

                elif (
                    started_copying
                    and "Account Number" in headers
                ):

                    break

            except:
                pass

        # ---------------------------------------------------
        # PARAGRAPHS
        # ---------------------------------------------------

        elif isinstance(block, Paragraph):

            if not started_copying:
                continue

            text = block.text.strip()

            if not text:
                continue

            # ---------------------------------------------------
            # STOP EXTRA CONTENT
            # ---------------------------------------------------

            if text.startswith(
                "Format for restoration"
            ):
                break

            # ---------------------------------------------------
            # ALWAYS COPY
            # ---------------------------------------------------

            always_copy = [

                "Chargeback Efforts:",
                "Finding of CLIC",
                "Recommendation with Justification:",
                "It is observed from SMS delivery report",
                "Had the customer called",
                "We confirm that",
                "Based on the above facts",
                "CM (FMC)",
                "Customer Liability Identification Cell",
                "LHO, Hyderabad",
                "Date:"
            ]

            copied = False

            for item in always_copy:

                if text.startswith(item):

                    para = output_doc.add_paragraph()

                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    para.paragraph_format.space_after = Pt(8)

                    run = para.add_run(text)

                    run.font.size = Pt(10)

                    if (
                        "Chargeback Efforts:" in text
                        or "Finding of CLIC" in text
                        or "Recommendation with Justification:" in text
                    ):

                        run.bold = True

                    copied = True
                    break

            if copied:
                continue

            # ---------------------------------------------------
            # CHANNEL DETECTION
            # ---------------------------------------------------

            if text.startswith(
                "ATM / POS / OTHPG:"
            ):

                inside_correct_channel = (
                    matched_channel in ["ATM", "POS"]
                )

            elif text.startswith(
                "UPI:"
            ):

                inside_correct_channel = (
                    matched_channel == "UPI"
                )

            elif text.startswith(
                "INB / YONO:"
            ):

                inside_correct_channel = (
                    matched_channel == "YONO"
                )

            elif text.startswith(
                "YONO Cash:"
            ):

                inside_correct_channel = False

            # ---------------------------------------------------
            # COPY CHANNEL CONTENT
            # ---------------------------------------------------

            if inside_correct_channel:

                para = output_doc.add_paragraph()

                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                para.paragraph_format.space_after = Pt(8)

                run = para.add_run(text)

                run.font.size = Pt(10)

                if (
                    text.startswith("ATM / POS / OTHPG:")
                    or text.startswith("UPI:")
                    or text.startswith("INB / YONO:")
                ):

                    run.bold = True

    # ---------------------------------------------------
    # SAVE FILE
    # ---------------------------------------------------

    output_path = f"outputs/{account_number}.docx"

    output_doc.save(output_path)

    return output_path